"""RENDER stage: pure Python .docx writing with built-in Word styles."""

from __future__ import annotations

from pathlib import Path

from docx import Document


def render_docx(title: str, sections: list[dict], output_path: Path) -> None:
    """Write title and heading/body sections to a Word document."""

    doc = Document()
    doc.add_heading(title, level=0)

    for section in sections:
        doc.add_heading(section["heading"], level=1)
        for paragraph in section["body"].split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
